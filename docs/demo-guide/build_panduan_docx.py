#!/usr/bin/env python3
"""Generate Panduan Demo E2E Collabite DOCX from narrate() scene screenshots.

Sumber kebenaran gambar = docs/demo-guide/assets/scenes/manifest.json
yang dihasilkan saat:
  DEMO_CAPTURE_SCENES=1 DEMO_HEADLESS=1 DEMO_STEP_MS=400 DEMO_SLOWMO=0 npm run test:e2e:demo
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parents[2]
ASSETS = Path(__file__).resolve().parent / "assets"
SCENES = ASSETS / "scenes"
OUT = ROOT / "docs" / "PANDUAN_DEMO_E2E_COLLABITE.docx"
FLOW = ASSETS / "flow_demo.png"
MANIFEST = SCENES / "manifest.json"


def set_run_font(run, size=11, bold=False, color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, size=16 if level == 1 else 13, bold=True)


def add_para(doc: Document, text: str, *, bold: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(6)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=11)


def add_image(doc: Document, path: Path, width_cm: float = 15.2, caption: str | None = None) -> None:
    if not path.exists() or path.stat().st_size < 5_000:
        add_para(doc, f"[Gambar tidak tersedia: {path.name}]", bold=True, size=10)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        set_run_font(r, size=9, color=RGBColor(0x55, 0x55, 0x55))


def draw_flow(entries: list[dict]) -> None:
    """Flow dari scene unik (scene label) dalam urutan manifest."""
    seen: list[str] = []
    for e in entries:
        s = e["scene"]
        if not seen or seen[-1] != s:
            seen.append(s)

    colors = {
        "PEMBUKA": "#18181B",
        "PENUTUP": "#18181B",
        "UMKM": "#0063D1",
        "CREATOR": "#E11D48",
        "ADMIN": "#16A34A",
        "WORKSPACE": "#7C3AED",
        "DEAL": "#7C3AED",
        "UNDANG": "#0063D1",
        "SELESAI": "#7C3AED",
    }

    def color_for(label: str) -> str:
        u = label.upper()
        for key, col in colors.items():
            if key in u:
                return col
        return "#52525B"

    n = max(len(seen), 1)
    w, box_h, gap = 1200, 64, 18
    h = 80 + n * (box_h + gap) + 80
    img = Image.new("RGB", (w, h), "#FFFDF8")
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 20)
        font_sm = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 15)
    except OSError:
        font = ImageFont.load_default()
        font_sm = font

    draw.text((40, 20), "Alur Demo (urutan scene dari rekaman nyata)", fill="#18181B", font=font)
    x0, box_w, y = 260, 680, 60
    for i, label in enumerate(seen):
        col = color_for(label)
        draw.rectangle([x0 + 3, y + 3, x0 + box_w + 3, y + box_h + 3], fill="#18181B")
        draw.rectangle([x0, y, x0 + box_w, y + box_h], outline="#18181B", width=3, fill="#FFFFFF")
        draw.rectangle([x0, y, x0 + 12, y + box_h], fill=col)
        draw.text((x0 + 24, y + 20), label, fill=col, font=font_sm)
        if i < len(seen) - 1:
            cy = y + box_h
            draw.line([(x0 + box_w // 2, cy), (x0 + box_w // 2, cy + gap)], fill="#18181B", width=3)
            ax, ay = x0 + box_w // 2, cy + gap
            draw.polygon([(ax, ay), (ax - 7, ay - 10), (ax + 7, ay - 10)], fill="#18181B")
        y += box_h + gap

    img.save(FLOW, "PNG")


def build() -> Path:
    if not MANIFEST.exists():
        raise SystemExit(
            "manifest.json belum ada. Jalankan dulu:\n"
            "  DEMO_CAPTURE_SCENES=1 DEMO_HEADLESS=1 DEMO_STEP_MS=400 DEMO_SLOWMO=0 npm run test:e2e:demo"
        )

    entries = json.loads(MANIFEST.read_text())
    if not entries:
        raise SystemExit("manifest.json kosong")

    ASSETS.mkdir(parents=True, exist_ok=True)
    draw_flow(entries)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("PANDUAN DEMO UI COLLABITE")
    set_run_font(r, size=22, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "Alur penuh npm run test:e2e:demo (Opsi A)\n"
        "Gambar = screenshot tepat saat banner narasi tiap langkah"
    )
    set_run_font(r, size=12, color=RGBColor(0x3F, 0x3F, 0x46))

    add_heading(doc, "1. Cara menjalankan", 1)
    add_bullets(
        doc,
        [
            "Full: npm run test:e2e:demo",
            "Pendek (tanpa undangan): npm run test:e2e:demo:short",
            "Tempo presentasi: DEMO_STEP_MS=3000 DEMO_SLOWMO=700 npm run test:e2e:demo",
            "Capture ulang gambar panduan: DEMO_CAPTURE_SCENES=1 DEMO_HEADLESS=1 DEMO_STEP_MS=400 DEMO_SLOWMO=0 npm run test:e2e:demo",
            "Lalu regenerate DOCX: /tmp/collabite-docx-venv/bin/python docs/demo-guide/build_panduan_docx.py",
        ],
    )

    add_heading(doc, "2. Diagram alur (dari scene nyata)", 1)
    add_image(doc, FLOW, width_cm=12.5, caption="Gambar 1. Urutan scene sesuai rekaman demo")

    add_heading(doc, "3. Langkah demi langkah (gambar = info)", 1)
    add_para(
        doc,
        f"Total {len(entries)} langkah. Setiap gambar diambil saat banner narasi "
        "tampil di layar — judul & catatan di bawah sama dengan teks banner.",
    )

    for e in entries:
        idx = e["index"]
        add_heading(doc, f"3.{idx} {e['scene']} — {e['title']}", 2)
        if e.get("note"):
            add_para(doc, f"Catatan: {e['note']}", size=10)
        add_image(
            doc,
            SCENES / e["file"],
            width_cm=15.2,
            caption=f"Gambar 3.{idx}. {e['scene']} — {e['title']}",
        )

    add_heading(doc, "4. Mapping file sumber", 1)
    add_bullets(
        doc,
        [
            "Orkestrator: tests/E2E/demo/demo-full-flow.spec.ts",
            "Helper narasi + capture: tests/E2E/demo/demo-helpers.ts",
            "Modul babak: tests/E2E/demo/modules/",
            "Config: playwright.demo.config.ts",
            "Manifest gambar: docs/demo-guide/assets/scenes/manifest.json",
        ],
    )

    add_heading(doc, "5. Tips", 1)
    add_bullets(
        doc,
        [
            "Jangan pakai frame video mentah dengan timestamp tebak-tebakan — mudah mismatch.",
            "Selalu regenerate dari DEMO_CAPTURE_SCENES=1 agar gambar = informasi.",
            "Presentasi: putar video.webm atau buka DOCX ini; pause di tiap banner BABAK.",
        ],
    )

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path} ({path.stat().st_size} bytes) from {len(json.loads(MANIFEST.read_text()))} scenes")
